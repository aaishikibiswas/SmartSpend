from __future__ import annotations

import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Iterable

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches
from sklearn.metrics import confusion_matrix


MODEL_PATTERN = re.compile(r"model_v(\d+)\.pkl$")


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def configure_logger(results_dir: Path, log_name: str) -> logging.Logger:
    results_dir.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger(log_name)
    logger.setLevel(logging.INFO)
    logger.handlers.clear()

    fmt = logging.Formatter("%(asctime)s | %(levelname)s | %(message)s")

    f_handler = logging.FileHandler(results_dir / "train.log", encoding="utf-8")
    f_handler.setFormatter(fmt)
    logger.addHandler(f_handler)

    s_handler = logging.StreamHandler()
    s_handler.setFormatter(fmt)
    logger.addHandler(s_handler)
    return logger


def append_metrics_csv(results_dir: Path, row: Dict[str, object]) -> Path:
    results_dir.mkdir(parents=True, exist_ok=True)
    metrics_path = results_dir / "metrics.csv"
    row_df = pd.DataFrame([row])
    if metrics_path.exists():
        existing = pd.read_csv(metrics_path)
        out = pd.concat([existing, row_df], ignore_index=True)
        out.to_csv(metrics_path, index=False)
    else:
        row_df.to_csv(metrics_path, index=False)
    return metrics_path


def next_model_version(models_dir: Path) -> int:
    models_dir.mkdir(parents=True, exist_ok=True)
    versions = []
    for model_path in models_dir.glob("model_v*.pkl"):
        match = MODEL_PATTERN.search(model_path.name)
        if match:
            versions.append(int(match.group(1)))
    return (max(versions) + 1) if versions else 1


def latest_model_path(models_dir: Path) -> Path | None:
    found = []
    for model_path in models_dir.glob("model_v*.pkl"):
        match = MODEL_PATTERN.search(model_path.name)
        if match:
            found.append((int(match.group(1)), model_path))
    if not found:
        return None
    return sorted(found, key=lambda x: x[0])[-1][1]


def save_confusion_matrix_plot(y_true, y_pred, out_path: Path, title: str) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    cm = confusion_matrix(y_true, y_pred)
    plt.figure(figsize=(6, 5))
    sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", cbar=False)
    plt.title(title)
    plt.xlabel("Predicted")
    plt.ylabel("Actual")
    plt.tight_layout()
    plt.savefig(out_path, dpi=180)
    plt.close()


def save_model_comparison_plot(comparison_df: pd.DataFrame, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    metric_cols = ["accuracy", "precision", "recall", "f1_score"]
    plot_df = comparison_df[["model_name"] + metric_cols].copy()
    melt_df = plot_df.melt(id_vars="model_name", value_vars=metric_cols, var_name="metric", value_name="score")

    plt.figure(figsize=(8, 5))
    sns.barplot(data=melt_df, x="metric", y="score", hue="model_name")
    plt.ylim(0, 1.05)
    plt.title("Model Comparison (Accuracy / Precision / Recall / F1)")
    plt.tight_layout()
    plt.savefig(out_path, dpi=180)
    plt.close()


def save_feature_importance_plot(
    feature_names: Iterable[str],
    importances: Iterable[float],
    out_path: Path,
    top_k: int = 12,
) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fi_df = pd.DataFrame({"feature": list(feature_names), "importance": list(importances)})
    fi_df = fi_df.sort_values("importance", ascending=False).head(top_k)

    plt.figure(figsize=(9, 5))
    sns.barplot(data=fi_df, x="importance", y="feature")
    plt.title(f"Top {top_k} Feature Importances")
    plt.tight_layout()
    plt.savefig(out_path, dpi=180)
    plt.close()


def generate_research_docx(
    reports_dir: Path,
    results_dir: Path,
    comparison_df: pd.DataFrame,
    best_model_name: str,
    best_reason: str,
) -> Path:
    reports_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("SmartSpend ML Analysis", level=0)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This study presents SmartSpend, a hybrid AI/ML financial analytics platform with a reproducible "
        "machine-learning pipeline and MLOps automation. The pipeline integrates transaction preprocessing, "
        "behavior-oriented feature engineering, supervised model comparison, and automated reporting. "
        "The system is production-ready with CI/CD retraining support and real-time communication integrations."
    )

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "SmartSpend is designed to support financial behavior prediction and decision support in near real-time. "
        "The architecture combines a backend API service, ML microservice workflows, and event-driven channels "
        "(WebSockets + Server-Sent Events) for live dashboard updates."
    )

    doc.add_heading("Methodology", level=1)
    doc.add_paragraph("Data preprocessing: schema validation, missing-value handling, date normalization, and type coercion.")
    doc.add_paragraph(
        "Feature engineering: monthly spending, category frequency, weekend ratio, anomaly indicators, "
        "income-expense ratio, and dominant category."
    )
    doc.add_paragraph("Model selection: RandomForest and LogisticRegression compared using F1-score as primary criterion.")

    doc.add_heading("Experimental Setup", level=1)
    doc.add_paragraph("Train/test split: 80/20 with fixed random seed for reproducibility.")
    doc.add_paragraph("Metrics: accuracy, precision, recall, F1-score.")
    doc.add_paragraph("Automation: GitHub Actions on push to main + daily cron retraining.")

    doc.add_heading("Results", level=1)
    doc.add_paragraph("Model comparison table:")
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "Accuracy"
    hdr[2].text = "Precision"
    hdr[3].text = "Recall"
    hdr[4].text = "F1-score"
    for _, row in comparison_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["model_name"])
        cells[1].text = f"{row['accuracy']:.4f}"
        cells[2].text = f"{row['precision']:.4f}"
        cells[3].text = f"{row['recall']:.4f}"
        cells[4].text = f"{row['f1_score']:.4f}"

    cm_png = results_dir / "confusion_matrix.png"
    comp_png = results_dir / "model_comparison.png"
    fi_png = results_dir / "feature_importance.png"

    for image_path, heading in [
        (cm_png, "Confusion Matrix"),
        (comp_png, "Model Comparison"),
        (fi_png, "Feature Importance"),
    ]:
        if image_path.exists():
            doc.add_heading(heading, level=2)
            doc.add_picture(str(image_path), width=Inches(6.2))

    doc.add_heading("Analysis", level=1)
    doc.add_paragraph(
        "Confusion matrix interpretation indicates class-level prediction quality and misclassification behavior. "
        "Feature importance highlights dominant financial drivers influencing risk classification."
    )
    doc.add_paragraph(
        f"Best model selected: {best_model_name}. Selection rationale: {best_reason}"
    )

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "SmartSpend demonstrates a research-grade yet deployment-friendly ML workflow with reproducible experiments, "
        "automated retraining, and extensible MLOps controls suitable for academic and applied fintech settings."
    )

    out_path = reports_dir / "SmartSpend_ML_Analysis.docx"
    doc.save(out_path)
    return out_path

